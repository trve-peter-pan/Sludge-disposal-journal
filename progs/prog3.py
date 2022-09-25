import shutil, os, time, os.path
from collections import defaultdict

import openpyxl
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_BREAK
from konfs.konfs import rarpath


def prog3(fullpathtoxlsxfile):
    """
    Создание актов в Word из файла Excel + создание архива rar
    """
    rusmon={1:"января", 2:"февраля", 3:"марта", 4:"апреля", 5:"мая", 6:"июня",
            7:"июля", 8:"августа", 9:"сентября", 10:"октября", 11:"ноября", 12:"декабря"}

    # TODO: осуществить создание списка ↓ при открытии GUI приложения
    wbFIO = openpyxl.load_workbook(r"..\konfs\Шаблоны\ФИО.xlsx")
    wsFI0 = wbFIO.worksheets[0]
    lastrow = 0
    for r in range(1, wsFI0.max_row+1):
        if wsFI0.cell(row=r, column=1).value is None:
            lastrow = r
            break
    FIO_MPS = {wsFI0.cell(row=i, column=2).value: {wsFI0.cell(row=1, column=j).value: wsFI0.cell(row=i, column=j).value
                                                   for j in range(3, wsFI0.max_column+1)} for i in range(2, lastrow)}
    start_time = time.time()
    actdir=r"..\2_Журнал_и_акты\Акты_НШЛ"
    wb = openpyxl.load_workbook(fullpathtoxlsxfile)
    ws0 = wb.worksheets[0]
    if os.path.exists(r"..\3_Архив"):
        shutil.rmtree(r"..\3_Архив")
    shutil.rmtree(actdir, ignore_errors=True)
    os.makedirs(actdir)
    os.makedirs(r"..\3_Архив")
    lastrow = 0
    for r in range(10, ws0.max_row+1):
        if ws0.cell(row=r, column=1).value is None:
            lastrow = r
            break
    table_for_act=[[ws0.cell(row=i, column=j).value for j in (1,2,3,4,5,7)] for i in range(10, lastrow)]
    wb.close()
    daterange = "(ошибка)"
    if table_for_act:
        daterange = f'({(table_for_act[0][1]).strftime("%d.%m.%y")}-{(table_for_act[-1][1]).strftime("%d.%m.%y")})'

    listmps=[table_for_act[i][2] for i in range(0, len(table_for_act))]
    # print("listmps=", listmps)
    tally = defaultdict(list)
    for i, item in enumerate(listmps):
        tally[item].append(i)
        # print(i, item)
    a=list(tally.items())
    docum = DocxTemplate(r"..\konfs\Шаблоны\Шаблон.docx")
    # print(len(table_for_act))
    for i in range(0, len(a)):
        for j in a[i][1]:
            template = Document(r"..\konfs\Шаблоны\Шаблон.docx")
            # print(j)
            if a[i][0] == 'Станция не определена':
                print(f"ВНИМАНИЕ! Строки {a[i][1]} в журнале НШЛ - не определена станция")
                break
            context = { 'НПС': FIO_MPS[a[i][0]]["НПС"],
                        'начальник_НПС': FIO_MPS[a[i][0]]["Начальник НПС"],
                        'должность_1': "начальник ЛАЭС",
                        'должность_2': "зам. начальника ЛАЭС",
                        'ФИО_1': FIO_MPS[a[i][0]]["Начальник ЛАЭС"],
                        'ФИО_2': FIO_MPS[a[i][0]]["Зам. начальника ЛАЭС"],
                        'номер': table_for_act[j][5], 'МН': table_for_act[j][3],
                        'масса': str(table_for_act[j][4]).replace(".", ","),
                        'день': f"{(table_for_act[j][1]).date().day:02d}",
                        'месяц': rusmon[(table_for_act[j][1]).date().month],
                        'год': (table_for_act[j][1]).date().year}
            # print(context)
            docum.render(context)
            print(f"В файл {(table_for_act[-1][1]).date().month:02d}_"
                  f"Акты_{a[i][0]}_{(table_for_act[-1][1]).date().year}.docx добавлен акт №{table_for_act[j][5]} "
                  f"от {table_for_act[j][1].date()}")
            docum.save(f"{actdir}\\{(table_for_act[-1][1]).date().month:02d}_"
                       f"Акты_{a[i][0]}_{(table_for_act[-1][1]).date().year}.docx")
            # print(a[i][1].index(j), len(a[i][1]) - 1)
            if a[i][1].index(j) < len(a[i][1]) - 1:
                last_paragraph = docum.paragraphs[-1]
                run = last_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                for element in template.element.body:
                    docum.element.body.append(element)
            docum.save(f"{actdir}\\{(table_for_act[-1][1]).date().month:02d}_"
                       f"Акты_{a[i][0]}_{(table_for_act[-1][1]).date().year}.docx")
            docum = DocxTemplate(f"{actdir}\\{(table_for_act[-1][1]).date().month:02d}_"
                                 f"Акты_{a[i][0]}_{(table_for_act[-1][1]).date().year}.docx")
        docum = DocxTemplate(r"..\konfs\Шаблоны\Шаблон.docx")
    print("Затрачено времени: %s секунд " % (time.time() - start_time))
    path_to_acts=os.path.abspath(f"{actdir}")
    archive_name=f"..\\3_Архив\\{daterange}.rar"
    root_dir=f"..\\2_Журнал_и_акты"
    rar = rarpath
    os.system(f'{rar} a {archive_name} {root_dir}')
    return path_to_acts, os.path.abspath(archive_name)

