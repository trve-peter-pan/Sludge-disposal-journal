import re, os, time
from os.path import join
from collections import defaultdict
from datetime import datetime
from decimal import Decimal
from konfs.konfs import MPS, PatternDict, SH, SS, ss

import openpyxl
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment, NamedStyle, colors
from docx import Document


def prog1(folder1):
    # перечень файлов в указанной папке в виде простого списка полных путей к этим файлам
    def ListFromFolder(folder, fileext):
        list_folder_files = []
        for root, dirs, files in os.walk(folder):
            list_folder_files.extend([join(root, file) for file in files if file.endswith(fileext) and "~" not in file])
            dirs.clear()
        return list_folder_files

    # таблица ДТВ в виде списка без обработки из указанного файла
    def raw_DTV_table_from_docx(fullpathtodocxfile):
        docxname = os.path.basename(fullpathtodocxfile)
        rawDTVtable = []
        if "docx" in docxname:
            m = None
            doc = Document(fullpathtodocxfile)
            properties = doc.core_properties
            all_tables = doc.tables
            data_tables = {i: None for i in range(len(all_tables))}
            for i, table in enumerate(all_tables):
                data_tables[i] = [[] for _ in range(len(table.rows))]
                for j, row in enumerate(table.rows):
                    for cell in row.cells:
                        data_tables[i][j].append(cell.text)
                        if "Наименование технологического участка" in data_tables[i][j] or "(результаты извлечения)" in \
                                data_tables[i][j]:
                            m = i
            rawDTVtable = data_tables[m]
            # добавим сведения о файле в таблицу
            docxname = str(os.path.basename(fullpathtodocxfile)).replace(SS, ss)
            prop_list = [docxname, properties.author, properties.last_modified_by, properties.created,
                         properties.modified,
                         properties.last_printed, fullpathtodocxfile]
            if len(rawDTVtable) > 0:
                for i in range(0, len(rawDTVtable)):
                    rawDTVtable[i].extend(prop_list)
        return rawDTVtable

    # список индексов повторяющихся элементов в указанном невложенном простом списке
    def list_duplicates(seq):
        y = defaultdict(list)
        for i, item in enumerate(seq):
            y[item].append(i)
        x = ((key, locs) for key, locs in y.items() if len(locs) > 1)
        b = []
        for dup in x:
            b.append([list((dup)[1])[i] for i in range(1, len(list((dup)[1])))])
        c = [b[i][j] for i in range(0, len(b)) for j in range(0, len(b[i]))]
        return c

    start_time = time.time()
    DTVtable = []
    for n, val in enumerate(ListFromFolder(folder1, "docx"), 1):
        print(n, os.path.basename(val))
        DTVtable.extend(raw_DTV_table_from_docx(val))
    # Удаляем лишние строки и столбцы по условию и создаем дубликат столбца
    DTVtable = [[DTVtable[i][j] for j in (0, 3, 4, 4, 5, 6, 7, 8, 9, 10, 11)] for i in range(0, (len(DTVtable)))]
    # фильтр регуляркой, оставляем только сочетания:
    for i in range(0, len(DTVtable)):
        result = re.findall(r"(\wзвле.*(\d){1,6}.{1,4}( л|л |л.ж|кг|тв|жид))", DTVtable[i][2], flags=re.IGNORECASE)
        if result:
            DTVtable[i][2] = re.sub('\s(?=\s)|', "", (result[0][0]).strip())
        else:
            DTVtable[i][2] = 0
    DTVtable = [[DTVtable[i][j] for j in range(0, (len(DTVtable[i])))] for i in range(0, (len(DTVtable))) if
                DTVtable[i][2] != 0]

    # Удаляем дубликаты строк по столбцу с извлечением (3)
    for i in range(0, len(DTVtable)):
        DTVtable[i][3] = str((re.sub('\\n', " ", str(DTVtable[i][3]))).strip())
    duplicates = list_duplicates([DTVtable[i][3] for i in range(0, (len(DTVtable)))])
    DTVtable = [[DTVtable[i][j] for j in range(0, len(DTVtable[i]))] for i in range(0, len(DTVtable)) if
                i not in duplicates]

    # Добавляем столбцы с МН и участком
    for i in range(0, (len(DTVtable))):
        DTVtable[i].insert(0, "Не определено наименование МН")
        DTVtable[i].insert(1, "Не определено наименование участка МН")
        DTVtable[i][2] = re.sub("\s", " ", DTVtable[i][2], flags=re.MULTILINE)
        for MN_pattern in PatternDict["MN_pattern"]:
            result = re.match(MN_pattern, DTVtable[i][2], flags=re.IGNORECASE)
            if result and DTVtable[i][0] == 'Не определено наименование МН':
                DTVtable[i][0] = PatternDict["MN_pattern"][MN_pattern]
                Sub_MN_pattern = PatternDict["Sub_MN_pattern"][DTVtable[i][0]]
                if type(Sub_MN_pattern) is str:
                    DTVtable[i][1] = Sub_MN_pattern
                else:
                    for key in Sub_MN_pattern:
                        result = re.match(key, DTVtable[i][2], flags=re.IGNORECASE)
                        if result and DTVtable[i][1] == "Не определено наименование участка МН":
                            DTVtable[i][1] = Sub_MN_pattern[key]

    # Добавляем столбец:
    for i in range(0, (len(DTVtable))):
        DTVtable[i].insert(0, "Станция не определена")
        for j in MPS:
            if DTVtable[i][2] in j:
                mps_name = j[0]
                DTVtable[i][0] = mps_name

    # обработка столбца 4 с работником (удаление повторяющихся значений и лишних пробелов, переносов строк)
    for i in range(0, len(DTVtable)):
        DTVtable[i][4] = str((re.sub(r'\b([^_]+)(\s+\1)+\b', r"\1", (re.sub('\\n', " ", str(DTVtable[i][4])))).strip()))

    # находим даты в столбце с извлечением и вносим отдельными столбцами в таблицу
    for i in range(0, (len(DTVtable))):
        DTVtable[i].insert(7, "Не определено")
        DTVtable[i].insert(7, "Время приема неизвестно")
        DTVtable[i].insert(7, "Время пуска неизвестно")
        DTVtable[i].insert(7, "Кол-во жидкого не определено")
        DTVtable[i].insert(7, "Кол-во твердого не определено")
        DTVtable[i].insert(7, datetime.strptime("00:00 01.01.01", "%H:%M %d.%m.%y"))
        result1 = re.findall("(\d\d:\d\d\s\d\d\.\d\d\.\d\d(?=.{0,7}пуск))", DTVtable[i][6], flags=re.IGNORECASE)
        if result1 and DTVtable[i][10] == 'Время пуска неизвестно':
            DTVtable[i][10] = datetime.strptime((result1[0]).strip(), "%H:%M %d.%m.%y")
        result2 = re.findall("(\d\d:\d\d\s\d\d\.\d\d\.\d\d(?=.{0,7}камер))", DTVtable[i][6], flags=re.IGNORECASE)
        if result2 and DTVtable[i][11] == 'Время приема неизвестно':
            DTVtable[i][11] = datetime.strptime((result2[0]).strip(), "%H:%M %d.%m.%y")
        if result1 and result2:
            DTVtable[i][12] = f'{str(divmod((DTVtable[i][11] - DTVtable[i][10]).total_seconds(), 3600)[0])} ч, ' \
                              f'{str(divmod(divmod((DTVtable[i][11] - DTVtable[i][10]).total_seconds(), 3600)[1], 60)[0])} мин.'
        # TODO: Нужна функция преобразования любой даты с лишними символами в дататайм
        result = re.findall("((\w\d{1,3}\s{0,1}[-\s:\.]\s{0,1}\d{1,3})\s{0,1}(\d{1,3}\s{0,1}[-\s:\.\\/]\s{0,1}\d{1,3}"
                            "\s{0,1}[-\s:\.\\/]\s{0,1}\d{1,3}))(?!.*((.\d{1,3}\s{0,1}[-\s:\.]\s{0,1}\d{1,3})\s{0,1}(\d{1,3}"
                            "\s{0,1}[-\s:\.\\/]\s{0,1}\d{1,3}\s{0,1}[-\s:\.\\/]"
                            "\s{0,1}\d{1,3})))(?=.*звле)", DTVtable[i][6], flags=re.IGNORECASE)
        if result and DTVtable[i][7] == datetime.strptime("00:00 01.01.01", "%H:%M %d.%m.%y"):
            dataresult = re.sub("[А-Яа-я]|\s(?=:)|(?<=:)\s|\s(?=\.)|(?<=\.)\s|\s(?=/)|(?<=/)\s|\s(?=-)|"
                                "(?<=-)\s|(?<=\d\d)\d(?=\s)|(?<=\s)\d(?=\d\d)|^\s|$\s", "", str(result[0][0]))
            a = re.match("^\d\d:\d\d \d\d\.\d\d\.\d\d$", dataresult)
            if a:
                b = datetime.strptime(str(dataresult), "%H:%M %d.%m.%y")
                DTVtable[i][7] = b

        result = re.findall(r"[\d,\.]{1,6}(?=[\D\s]{0,3}кг)(?![\D\s]{0,5}жид)|[\d,\.]{1,6}(?=[\D\s]{0,3}тв)",
                            DTVtable[i][6], flags=re.IGNORECASE)
        if result and DTVtable[i][8] == 'Кол-во твердого не определено':
            DTVtable[i][8] = sum(Decimal(x) for x in [str(x).replace(",", ".") for x in result])
        # Удалим двойные запятые и точки на цифрах, чтобы не было ошибки при преобразовании str во float
        DTVtable[i][6] = re.sub("[\.,](?=\d{1,5}[\.,])|(?<=\s),(?=\d)|(?<=\s)\.(?=\d)", " ", DTVtable[i][6],
                                flags=re.IGNORECASE)
        result = re.findall(r"[\d,\.]{1,6}(?=[\D\s]{0,3}кг)(?=[\D\s]{0,5}жид)|[\d,\.]{1,6}(?=[\D\s]{0,3}жид)|"
                            r"[\d,\.]{1,6}(?=[\D\s]{0,3}л)", DTVtable[i][6], flags=re.IGNORECASE)
        if result and DTVtable[i][9] == 'Кол-во жидкого не определено':
            DTVtable[i][9] = sum(Decimal(x) for x in [str(x).replace(",", ".") for x in result])
    # меняем порядок расположения столбцов:

    DTVtable = [[DTVtable[i][j] for j in (0, 2, 4, 7, 8, 6, 13, 15, 16, 17, 18, 9, 12, 3, 19)] for i in
                range(0, (len(DTVtable)))]
    # ошибка при х не равном типу дататайм, нужен более четкий фильтр
    DTVtable.sort(key=lambda x: x[3])
    # добавляем столбец с нумерацией
    for i in range(0, (len(DTVtable))):
        DTVtable[i].insert(0, i + 1)
    # cоздаем таблицу для журнала
    jtable = [[DTVtable[i][j] for j in (4, 1, 2, 5)] for i in range(0, (len(DTVtable)))]
    # print(jtable)
    for i in range(0, (len(jtable))):
        jtable[i].insert(1, (jtable[i][0]).date())

    jtable = [[jtable[i][j] for j in (1, 2, 3, 4)] for i in range(0, (len(jtable)))]

    key = lambda item: (item[0], item[1], item[2])
    keys = set(map(key, jtable))
    jtable_out = [[k[0], k[1], k[2], sum(Decimal(n[3]) for n in jtable if k == key(n))] for k in keys]
    jtable = jtable_out

    # сортируем по возрастанию даты
    jtable.sort(key=lambda x: x[0])

    # добавляем столбец с нумерацией
    m = str(jtable[-1][0].month)

    for i in range(0, (len(jtable))):
        jtable[i].insert(0, i + 1)
        jtable[i].insert(-1, f'{m}-{str(jtable[i][0])}')
        jtable[i].insert(-1, SH)
    for i in range(0, (len(jtable))):
        jtable[i][-1] = (jtable[i][-1] / 1000)
    jtable = [[jtable[i][j] for j in (0, 1, 2, 3, 6, 6, 4, 1, 5, 6, 6)] for i in range(0, (len(jtable))) if
              jtable[i][6] != 0]
    # добавляем столбец с нумерацией
    for i in range(0, (len(jtable))):
        jtable[i][0] = i + 1
    sum_jtable = sum([jtable[i][4] for i in range(0, (len(jtable)))])
    for i in range(0, (len(jtable))):
        jtable[i][6] = f'{m}-{str(jtable[i][0])}'

    Jtable_itog = [
        ["", "", "", "Итого за месяц:", sum_jtable, sum_jtable, "", "", "", sum_jtable, sum_jtable, "", "", "", "", "",
         "", "0,000", "0,000"]]

    wb = openpyxl.load_workbook(r"..\konfs\Шаблоны\Шаблон.xlsx")
    ws0 = wb.worksheets[0]
    ws = wb.worksheets[1]
    # wb.save("Шаблон.xlsx")
    for n, row in enumerate(jtable, 10):
        ws0.append(row)
        for i in (1, 5, 6, 7, 9, 10, 11):
            ws0.cell(row=n, column=i).style = "style9c"
        for i in (3, 4):
            ws0.cell(row=n, column=i).style = "style9"
        for i in (2, 8):
            ws0.cell(row=n, column=i).style = "datestyle_short"
        for i in range(1, 22):
            Thin = Side(border_style='thin', color="FF000000")
            ws0.cell(row=n, column=i).border = Border(top=Thin, bottom=Thin, left=Thin, right=Thin)

    lr = ws0.max_row
    for row in Jtable_itog:
        ws0.append(row)
        for i in (5, 6, 10, 11, 18, 19):
            ws0.cell(row=lr + 1, column=i).style = "style9cb"
        for i in (3, 4):
            ws0.cell(row=lr + 1, column=i).style = "style9rb"
        for i in range(1, 22):
            Thin = Side(border_style='thin', color="FF000000")
            ws0.cell(row=lr + 1, column=i).border = Border(top=Thin, bottom=Thin, left=Thin, right=Thin)

    for n, row in enumerate(DTVtable, 2):
        ws.append(row)
        for i in (1, 4, 14, 15):
            ws.cell(row=n, column=i).style = "style11"
        for i in (2, 3):
            ws.cell(row=n, column=i).style = "style11l"
        for i in (6, 13):
            ws.cell(row=n, column=i).style = "style16"
        for i in (7, 9):
            ws.cell(row=n, column=i).style = "style9"
        ws.cell(row=n, column=8).hyperlink = f".{DTVtable[n - 2][15]}"
        ws.cell(row=n, column=8).style = "style_9l_gl"
        for i in (5, 10, 11, 12):
            ws.cell(row=n, column=i).style = "datestyle"
        for i in range(1, 16):
            if divmod(n, 2)[1] == 0:
                ws.cell(row=n, column=i).fill = PatternFill('solid', fgColor="D9D9D9")
        for i in range(1, 16):
            badcell = ws.cell(row=n, column=i).value
            if badcell == 'Станция не определена' or badcell == datetime.strptime("00:00 01.01.01", "%H:%M %d.%m.%y") or \
                    badcell == 'Не определено наименование участка МН' or badcell == 'Кол-во твердого не определено':
                for j in range(1, 16):
                    ws.cell(row=n, column=j).fill = PatternFill('solid', fgColor="D97147")
        ws.delete_cols(16)
    tab = Table(displayName="Сводка", ref=f"A1:{ws.cell(row=ws.max_row, column=ws.max_column).coordinate}")

    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 35
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 10
    ws.column_dimensions['G'].width = 55
    ws.column_dimensions['H'].width = 15
    ws.column_dimensions['I'].width = 15
    ws.column_dimensions['J'].width = 15
    ws.column_dimensions['K'].width = 15
    ws.column_dimensions['L'].width = 15
    ws.column_dimensions['M'].width = 10
    ws.column_dimensions['N'].width = 17
    ws.column_dimensions['O'].width = 42

    style = TableStyleInfo(name="TableStyleLight9", showFirstColumn=False,
                           showLastColumn=False, showRowStripes=True, showColumnStripes=True)
    tab.tableStyleInfo = style
    # сохраняем чтобы перенести таблицу в новый файл из файла шаблона, чтобы не было ошибки задвоения имени таблицы
    wb.save(r"..\2_Журнал_и_акты\Журнал_НШЛ.xlsx")
    ws.add_table(tab)
    wb.save(r"..\2_Журнал_и_акты\Журнал_НШЛ.xlsx")

    print("Затрачено времени: %s секунд " % (time.time() - start_time))
    a = os.path.abspath(r"..\2_Журнал_и_акты\Журнал_НШЛ.xlsx")
    return a


if __name__ == '__main__':
    print(prog1(r'..\1_Сводки'))
